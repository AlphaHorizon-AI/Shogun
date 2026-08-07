import uuid
from copy import copy
from datetime import datetime
from types import SimpleNamespace

import openpyxl
import pytest
from docx import Document

from shogun.config import settings
from shogun.engine import flow_engine
from shogun.services.file_template import (
    TEMPLATE_MARKER,
    extract_file_template,
    format_template_guidance,
    parse_excel_rows,
    render_excel_template,
    render_word_template,
)


def test_excel_json_object_rows_follow_template_column_order():
    payload = """```json
[{"Amount": 20, "Item": "140006", "2026-07-01": 10},
 {"Item": "140023", "2026-07-01": 5, "Amount": 12}]
```"""

    rows = parse_excel_rows(
        payload,
        template_headers=[None, "Item", "Amount", datetime(2026, 7, 1)],
    )

    assert rows == [[None, "140006", 20, 10], [None, "140023", 12, 5]]


def test_excel_chunked_and_individually_serialized_json_rows_are_decoded():
    payload = """[Output from 'Extract & Map Data']:
```json
["{\\"Item\\":\\"A\\",\\"Amount\\":1}"]
```
```json
[{"Item":"B","Amount":2}]
```"""

    rows = parse_excel_rows(payload, template_headers=["Item", "Amount"])

    assert rows == [["A", 1], ["B", 2]]


def test_excel_incomplete_structured_json_is_rejected():
    with pytest.raises(ValueError, match="incomplete or truncated"):
        parse_excel_rows('```json\n[["A", "B"]')


def test_excel_manifest_and_adaptive_render_ignore_formatting_only_rows(tmp_path):
    source = tmp_path / "source.xlsx"
    output = tmp_path / "Output" / "result.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Plan"
    ws.append(["Description", "Item", "Quantity"])
    ws.append([None, "Nr.", "Stk."])
    ws.append(["Existing", "100", 1])
    ws["A100"].fill = openpyxl.styles.PatternFill(fill_type="solid", fgColor="FFFF00")
    wb.save(source)
    wb.close()

    payload = extract_file_template("source.xlsx", tmp_path)
    sheet = payload["manifest"]["sheets"][0]
    assert sheet["logical_range"] == "A1:C3"
    assert sheet["logical_columns"] == 3
    assert sheet["suggested_append_cell"] == "A4"

    changed = render_excel_template(
        source,
        output,
        '[["Generated", "200", 2]]',
        "append",
        "Plan",
        render_mode="adaptive",
    )
    assert changed == 3
    rendered = openpyxl.load_workbook(output)
    try:
        assert rendered["Plan"]["A3"].value == "Existing"
        assert rendered["Plan"]["A4"].value == "Generated"
        assert rendered["Plan"]["C4"].value == 2
        assert rendered["Plan"]["A100"].value is None
    finally:
        rendered.close()


def test_excel_template_matrix_width_is_validated():
    with pytest.raises(ValueError, match="exactly 3 values"):
        parse_excel_rows('[["A", "B"]]', template_headers=["One", "Two", "Three"])


def test_excel_missing_sheet_uses_only_template_sheet_with_warning(tmp_path, caplog):
    source = tmp_path / "source.xlsx"
    output = tmp_path / "result.xlsx"
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "Production Plan"
    worksheet.append(["Item", "Quantity"])
    workbook.save(source)
    workbook.close()

    caplog.set_level("WARNING", logger="shogun.file_template")
    changed = render_excel_template(
        source,
        output,
        '[["A", 2]]',
        "append",
        "Advanced Output",
        render_mode="adaptive",
    )

    assert changed == 2
    assert "using the template's only worksheet" in caplog.text
    rendered = openpyxl.load_workbook(output)
    try:
        assert rendered["Production Plan"]["A2"].value == "A"
        assert rendered["Production Plan"]["B2"].value == 2
    finally:
        rendered.close()


def test_excel_missing_sheet_rejects_ambiguous_multi_sheet_template(tmp_path):
    source = tmp_path / "source.xlsx"
    output = tmp_path / "result.xlsx"
    workbook = openpyxl.Workbook()
    workbook.active.title = "First"
    workbook.create_sheet("Second")
    workbook.save(source)
    workbook.close()

    with pytest.raises(ValueError, match="Available worksheets: 'First', 'Second'"):
        render_excel_template(
            source,
            output,
            '[["A"]]',
            "append",
            "Missing",
            render_mode="adaptive",
        )


def test_excel_markdown_summary_width_is_validated_against_template():
    payload = "| Item | Quantity |\n| --- | --- |\n| A | 2 |"

    with pytest.raises(ValueError, match="exactly 3 values"):
        parse_excel_rows(payload, template_headers=["Description", "Item", "Quantity"])


def test_excel_matrix_contract_rejects_markdown_even_when_width_matches():
    payload = "| Item | Quantity |\n| --- | --- |\n| A | 2 |"

    with pytest.raises(ValueError, match="two-dimensional JSON array"):
        parse_excel_rows(payload, require_structured_json=True)


def test_word_template_structure_and_one_shot_are_explicit(tmp_path):
    template = tmp_path / "report.docx"
    doc = Document()
    doc.add_heading("Quarterly Report", level=1)
    doc.add_paragraph("Customer: {{customer}}")
    doc.add_paragraph("Example conclusion")
    doc.save(template)

    structure = extract_file_template("report.docx", tmp_path, "structure_only")
    assert structure[TEMPLATE_MARKER] is True
    assert structure["format"] == "docx"
    assert "Quarterly Report" in structure["contract"]
    assert "customer" in structure["contract"]
    assert structure["example"] == ""

    one_shot = extract_file_template("report.docx", tmp_path, "one_shot", "replace")
    guidance = format_template_guidance(one_shot)
    assert "[POPULATED ONE-SHOT EXAMPLE]" in guidance
    assert "Example conclusion" in guidance
    assert "example records are non-authoritative" in guidance
    assert "reference-only output contract" in guidance
    assert "sole source of business records" in guidance


def test_excel_authoritative_baseline_contract_hides_business_rows(tmp_path):
    template = tmp_path / "master.xlsx"
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "Master"
    worksheet.append(["Description", "Item ID", "Quantity"])
    worksheet.append(["PRIVATE BASELINE VALUE", "A-100", 12])
    workbook.save(template)
    workbook.close()

    payload = extract_file_template(
        "master.xlsx",
        tmp_path,
        "baseline_merge",
        merge_key_columns="B",
        merge_preserve_columns="A",
    )
    guidance = format_template_guidance(payload)

    assert payload["authoritative_baseline"] is True
    assert payload["merge_strategy"] == "replace_matching_groups"
    assert payload["merge_key_columns"] == ["B"]
    assert payload["merge_preserve_columns"] == ["A"]
    assert payload["example"] == ""
    assert payload["manifest"]["sheets"][0]["preview_rows"] == [
        ["Description", "Item ID", "Quantity"]
    ]
    assert "PRIVATE BASELINE VALUE" not in guidance
    assert "authoritative downstream baseline" in guidance
    assert "Do not reproduce" in guidance


def test_excel_authoritative_baseline_replaces_entity_groups_and_preserves_lookup_columns(tmp_path):
    source = tmp_path / "master.xlsx"
    output = tmp_path / "Output" / "merged.xlsx"
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "Plan"
    worksheet.append(["Description", "Item ID", "Order ID", "Quantity", "Owner", "Comment"])
    worksheet.append(["Article A", "A", "O1", 1, "Alice", "Master note"])
    worksheet.append(["Article A", "A", "O2", 2, "Bob", "Old order"])
    worksheet.append(["Article B", "B", "O9", 9, "Carol", "Unrelated"])
    bold_font = copy(worksheet["A2"].font)
    bold_font.bold = True
    worksheet["A2"].font = bold_font
    workbook.save(source)
    workbook.close()

    changed = render_excel_template(
        source,
        output,
        """[["Article A", "A", "O1", 10, "", ""],
                 ["Article A", "A", "O3", 3, "", ""],
                 ["Article C", "C", "O7", 7, "Dana", "New"]]""",
        "replace",
        "Plan",
        render_mode="adaptive",
        guidance_mode="baseline_merge",
        merge_key_columns="Item ID",
        merge_preserve_columns="E:F",
    )

    assert changed == 24
    rendered = openpyxl.load_workbook(output)
    original = openpyxl.load_workbook(source)
    try:
        values = [
            [rendered["Plan"].cell(row, column).value for column in range(1, 7)]
            for row in range(2, 6)
        ]
        assert values == [
            ["Article A", "A", "O1", 10, "Alice", "Master note"],
            ["Article A", "A", "O3", 3, "Alice", "Master note"],
            ["Article B", "B", "O9", 9, "Carol", "Unrelated"],
            ["Article C", "C", "O7", 7, "Dana", "New"],
        ]
        assert rendered["Plan"]["A2"].font.bold is True
        assert original["Plan"]["D2"].value == 1
        assert original["Plan"]["A5"].value is None
    finally:
        rendered.close()
        original.close()


def test_excel_authoritative_baseline_rejects_runtime_rows_without_entity_key(tmp_path):
    source = tmp_path / "master.xlsx"
    output = tmp_path / "result.xlsx"
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.append(["Description", "Item ID", "Quantity"])
    worksheet.append(["Existing", "A", 1])
    workbook.save(source)
    workbook.close()

    with pytest.raises(ValueError, match="empty key"):
        render_excel_template(
            source,
            output,
            '[["Missing key", "", 2]]',
            "replace",
            guidance_mode="baseline_merge",
            merge_key_columns="B",
        )


def test_word_template_render_creates_copy_and_replaces_json_placeholders(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "Output" / "result.docx"
    doc = Document()
    doc.add_heading("Client Brief", level=1)
    doc.add_paragraph("Client: {{client}}")
    doc.add_paragraph("Summary: {{summary}}")
    doc.save(source)

    changed = render_word_template(
        source,
        output,
        '[Output from \'Samurai\']:\n{"client":"Acme","summary":"Ready"}',
        "replace",
    )

    assert changed == 2
    assert "{{client}}" in "\n".join(p.text for p in Document(source).paragraphs)
    rendered = "\n".join(p.text for p in Document(output).paragraphs)
    assert "Client: Acme" in rendered
    assert "Summary: Ready" in rendered


def test_word_adaptive_render_populates_repeating_table_rows(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "Output" / "result.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Quantity"
    table.cell(1, 0).text = "Example"
    table.cell(1, 1).text = "1"
    document.save(source)

    changed = render_word_template(
        source,
        output,
        '[["A", 2], ["B", 3]]',
        "replace",
        render_mode="adaptive",
    )

    rendered = Document(output)
    assert changed == 4
    assert [[cell.text for cell in row.cells] for row in rendered.tables[0].rows] == [
        ["Item", "Quantity"],
        ["A", "2"],
        ["B", "3"],
    ]


def test_excel_one_shot_replace_preserves_template_and_styles(tmp_path):
    source = tmp_path / "source.xlsx"
    output = tmp_path / "Output" / "result.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Orders"
    ws.append(["Customer", "Amount"])
    ws.append(["Example Co", 100])
    sample_font = copy(ws["A1"].font)
    sample_font.bold = True
    ws["A2"].font = sample_font
    wb.save(source)
    wb.close()

    payload = extract_file_template("source.xlsx", tmp_path, "one_shot", "replace")
    assert "Example Co" in payload["example"]
    assert "Customer | Amount" in payload["contract"]

    changed = render_excel_template(
        source,
        output,
        "| Customer | Amount |\n| --- | --- |\n| New Co | 250 |",
        "replace",
        "Orders",
    )
    assert changed == 2

    original = openpyxl.load_workbook(source)
    rendered = openpyxl.load_workbook(output)
    try:
        assert original["Orders"]["A2"].value == "Example Co"
        assert rendered["Orders"]["A2"].value == "New Co"
        assert rendered["Orders"]["B2"].value == "250"
        assert rendered["Orders"]["A2"].font.bold is True
    finally:
        original.close()
        rendered.close()


def test_excel_one_shot_replace_overrides_adaptive_placement(tmp_path):
    source = tmp_path / "source.xlsx"
    output = tmp_path / "Output" / "result.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Orders"
    ws.append(["Item", "Quantity"])
    ws.append(["EXAMPLE-A", 100])
    ws.append(["EXAMPLE-B", 200])
    wb.save(source)
    wb.close()

    changed = render_excel_template(
        source,
        output,
        '[["SOURCE-A", 12]]',
        "replace",
        "Orders",
        render_mode="adaptive",
    )

    assert changed == 2
    rendered = openpyxl.load_workbook(output)
    try:
        assert rendered["Orders"]["A2"].value == "SOURCE-A"
        assert rendered["Orders"]["B2"].value == 12
        assert rendered["Orders"]["A3"].value is None
        assert rendered["Orders"]["B3"].value is None
    finally:
        rendered.close()


def test_excel_one_shot_replace_preserves_sparse_secondary_header_row(tmp_path):
    source = tmp_path / "source.xlsx"
    output = tmp_path / "Output" / "result.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Plan"
    ws.append([None, "Artikel-Nr", "Fertigungsauftrag", "Stk.", "Month"])
    ws.append([None, None, "Nr.", "Qty", None])
    ws.append(["EXAMPLE", 100, 200, 1, None])
    wb.save(source)
    wb.close()

    changed = render_excel_template(
        source,
        output,
        '[["SOURCE", 140006, 20164555, 20, ""]]',
        "replace",
        "Plan",
        render_mode="adaptive",
    )

    assert changed == 5
    rendered = openpyxl.load_workbook(output)
    try:
        assert rendered["Plan"]["C2"].value == "Nr."
        assert rendered["Plan"]["D2"].value == "Qty"
        assert [rendered["Plan"].cell(3, column).value for column in range(1, 6)] == [
            "SOURCE",
            140006,
            20164555,
            20,
            None,
        ]
    finally:
        rendered.close()


def test_excel_anchored_replace_clears_remaining_example_rows(tmp_path):
    source = tmp_path / "source.xlsx"
    output = tmp_path / "Output" / "result.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Orders"
    ws.append(["Report title", None])
    ws.append(["Item", "Quantity"])
    ws.append([None, None])
    ws.append(["EXAMPLE-A", 100])
    ws.append(["EXAMPLE-B", 200])
    wb.save(source)
    wb.close()

    changed = render_excel_template(
        source,
        output,
        '[["SOURCE-A", 12]]',
        "replace",
        "Orders",
        start_cell="A4",
        render_mode="adaptive",
    )

    assert changed == 2
    rendered = openpyxl.load_workbook(output)
    try:
        assert rendered["Orders"]["A4"].value == "SOURCE-A"
        assert rendered["Orders"]["B4"].value == 12
        assert rendered["Orders"]["A5"].value is None
        assert rendered["Orders"]["B5"].value is None
    finally:
        rendered.close()


def test_template_ancestry_reaches_downstream_create_node():
    payload = {TEMPLATE_MARKER: True, "template_path": "Templates/report.docx"}
    predecessors = {"samurai": ["template"], "files": ["samurai"]}
    outputs = {"template": payload, "samurai": "generated content"}

    assert flow_engine._collect_upstream_file_templates("files", predecessors, outputs) == [payload]


def test_input_fan_out_triggers_template_and_parallel_branch_together():
    input_id = uuid.uuid4()
    template_id = uuid.uuid4()
    parallel_id = uuid.uuid4()
    samurai_id = uuid.uuid4()
    nodes = [
        SimpleNamespace(id=input_id),
        SimpleNamespace(id=template_id),
        SimpleNamespace(id=parallel_id),
        SimpleNamespace(id=samurai_id),
    ]
    edges = [
        SimpleNamespace(source_node_id=input_id, target_node_id=template_id),
        SimpleNamespace(source_node_id=input_id, target_node_id=parallel_id),
        SimpleNamespace(source_node_id=template_id, target_node_id=samurai_id),
    ]

    layers = flow_engine._topological_sort(nodes, edges)

    assert layers[0] == [str(input_id)]
    assert set(layers[1]) == {str(template_id), str(parallel_id)}
    assert layers[2] == [str(samurai_id)]


@pytest.mark.asyncio
async def test_word_create_uses_upstream_template_and_never_changes_source(tmp_path, monkeypatch):
    from shogun.office import config as office_config

    template = tmp_path / "Templates" / "brief.docx"
    template.parent.mkdir(parents=True)
    doc = Document()
    doc.add_paragraph("Customer: {{customer}}")
    doc.save(template)
    monkeypatch.setattr(settings, "workspace_path", tmp_path)
    monkeypatch.setattr(office_config, "load_office_config", lambda: SimpleNamespace(enabled=True))

    result = await flow_engine._exec_office(
        {"action": "word_create", "output_path": "Output", "output_filename": "brief.docx"},
        '{"customer":"Acme"}',
        run_id=uuid.uuid4(),
        trigger_type="manual",
        template_inputs=[
            {
                TEMPLATE_MARKER: True,
                "template_path": "Templates/brief.docx",
                "format": "docx",
                "example_handling": "replace",
            }
        ],
    )

    assert "Word document created" in result
    assert "{{customer}}" in Document(template).paragraphs[0].text
    assert "Customer: Acme" in Document(tmp_path / "Output" / "brief.docx").paragraphs[0].text


@pytest.mark.asyncio
async def test_excel_create_applies_upstream_authoritative_baseline_contract(tmp_path, monkeypatch):
    from shogun.office import config as office_config

    template = tmp_path / "Templates" / "master.xlsx"
    template.parent.mkdir(parents=True)
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "Plan"
    worksheet.append(["Item ID", "Order ID", "Quantity", "Owner"])
    worksheet.append(["A", "O1", 1, "Alice"])
    worksheet.append(["B", "O2", 2, "Bob"])
    workbook.save(template)
    workbook.close()
    monkeypatch.setattr(settings, "workspace_path", tmp_path)
    monkeypatch.setattr(office_config, "load_office_config", lambda: SimpleNamespace(enabled=True))

    template_payload = extract_file_template(
        "Templates/master.xlsx",
        tmp_path,
        "baseline_merge",
        merge_key_columns="A",
        merge_preserve_columns="D",
    )
    result = await flow_engine._exec_office(
        {
            "action": "excel_create",
            "output_path": "Output",
            "output_filename": "result.xlsx",
            "sheet_name": "Plan",
        },
        '[["A", "O1", 10, ""], ["C", "O3", 3, "Carol"]]',
        run_id=uuid.uuid4(),
        trigger_type="manual",
        template_inputs=[template_payload],
    )

    assert "Excel workbook created" in result
    rendered = openpyxl.load_workbook(tmp_path / "Output" / "result.xlsx")
    try:
        assert [
            [rendered["Plan"].cell(row, column).value for column in range(1, 5)]
            for row in range(2, 5)
        ] == [
            ["A", "O1", 10, "Alice"],
            ["B", "O2", 2, "Bob"],
            ["C", "O3", 3, "Carol"],
        ]
    finally:
        rendered.close()


def test_template_path_must_stay_in_workspace(tmp_path):
    outside = tmp_path.parent / "outside.docx"
    Document().save(outside)
    with pytest.raises(ValueError, match="workspace"):
        extract_file_template(str(outside), tmp_path)
