"""Template extraction and rendering for AgentFlow File Template nodes.

The template source is always treated as immutable. Rendering copies it to a
new output path before applying generated content.
"""

from __future__ import annotations

import json
import re
import shutil
from copy import copy
from pathlib import Path
from typing import Any

TEMPLATE_MARKER = "__shogun_file_template__"
SUPPORTED_TEMPLATE_SUFFIXES = {".docx", ".xlsx"}
_PLACEHOLDER_RE = re.compile(r"{{\s*([A-Za-z0-9_.-]+)\s*}}")


def resolve_workspace_template(template_path: str, workspace_root: Path) -> Path:
    """Resolve and validate a template path inside the configured workspace."""
    raw = str(template_path or "").strip().replace("\\", "/")
    if not raw:
        raise ValueError("Select a Word (.docx) or Excel (.xlsx) template file.")
    if ".." in Path(raw).parts:
        raise ValueError(f"Template path traversal blocked: {raw}")

    root = workspace_root.resolve()
    target = (root / raw).resolve()
    try:
        target.relative_to(root)
    except ValueError as exc:
        raise ValueError("Template file must remain inside the configured workspace.") from exc
    if target.suffix.lower() not in SUPPORTED_TEMPLATE_SUFFIXES:
        raise ValueError("File Template nodes currently support .docx and .xlsx files.")
    if not target.is_file():
        raise ValueError(f"Template file was not found: {raw}")
    return target


def extract_file_template(
    template_path: str,
    workspace_root: Path,
    guidance_mode: str = "structure_only",
    example_handling: str = "replace",
    max_chars: int = 12_000,
) -> dict[str, Any]:
    """Build a bounded template contract for the Samurai and Files node."""
    mode = guidance_mode if guidance_mode in {"structure_only", "one_shot"} else "structure_only"
    handling = example_handling if example_handling in {"replace", "append", "preserve"} else "replace"
    source = resolve_workspace_template(template_path, workspace_root)
    relative = source.relative_to(workspace_root.resolve()).as_posix()

    if source.suffix.lower() == ".docx":
        contract, example = _extract_word_contract(source)
        manifest = _word_template_manifest(source)
        file_format = "docx"
    else:
        contract, example = _extract_excel_contract(source)
        manifest = _excel_template_manifest(source)
        file_format = "xlsx"

    return {
        TEMPLATE_MARKER: True,
        "template_path": relative,
        "format": file_format,
        "guidance_mode": mode,
        "example_handling": handling,
        "contract": contract[:max_chars],
        "example": example[:max_chars] if mode == "one_shot" else "",
        "manifest": manifest,
    }


def format_template_guidance(payload: dict[str, Any]) -> str:
    """Format a template payload as a precise model-facing output contract."""
    mode = payload.get("guidance_mode", "structure_only")
    lines = [
        "[FILE TEMPLATE CONTRACT]",
        f"Template: {payload.get('template_path', '')}",
        f"Format: {payload.get('format', '')}",
        f"Guidance mode: {'one-shot example' if mode == 'one_shot' else 'structure only'}",
        f"Output handling: {payload.get('example_handling', 'replace')}",
        "Generate the requested content so it fits this exact template contract. Do not output the template itself.",
        "Return only the content/data that should be inserted into the new file.",
        "Treat the template as a reference-only output contract, never as a factual data source.",
        "Use non-template runtime inputs as the sole source of business records and values.",
        "Do not copy populated template values unless the same value is independently present in the runtime input.",
        "",
        str(payload.get("contract") or ""),
    ]
    if payload.get("manifest"):
        lines.extend(
            [
                "",
                "[MACHINE-READABLE TEMPLATE MANIFEST]",
                json.dumps(payload["manifest"], ensure_ascii=False, default=str),
            ]
        )
    if mode == "one_shot" and payload.get("example"):
        lines.extend(
            [
                "",
                "[POPULATED ONE-SHOT EXAMPLE]",
                "Use the following existing content only to learn layout, ordering, and formatting.",
                "These example records are non-authoritative. Do not reproduce them unless they are "
                "independently supported by the runtime input:",
                str(payload["example"]),
            ]
        )
    return "\n".join(lines).strip()


def render_word_template(
    source: Path,
    output: Path,
    context: str,
    handling: str,
    render_mode: str = "strict",
) -> int:
    """Copy a Word template and populate the copy. Returns changed item count."""
    from docx import Document

    _copy_template(source, output)
    doc = Document(str(output))
    content = _unwrap_flow_context(context)
    replacements = _context_replacements(content)
    changed = _replace_docx_placeholders(doc, replacements)

    if render_mode == "adaptive" and changed == 0:
        changed = _populate_first_word_table(doc, content, append=handling == "append")
        if changed == 0:
            for line in content.splitlines() or [content]:
                doc.add_paragraph(line)
                changed += 1
    elif handling == "append":
        for line in content.splitlines() or [content]:
            doc.add_paragraph(line)
            changed += 1
    elif handling == "preserve":
        if changed == 0:
            raise ValueError("Preserve unchanged requires at least one {{placeholder}} in the Word template.")
    elif changed == 0:
        changed = _replace_word_example_content(doc, content)

    doc.save(str(output))
    return changed


def render_excel_template(
    source: Path,
    output: Path,
    context: str,
    handling: str,
    sheet_name: str | None = None,
    start_cell: str | None = None,
    render_mode: str = "strict",
) -> int:
    """Copy an Excel template and populate the copy. Returns written cell count."""
    import openpyxl

    _copy_template(source, output)
    wb = openpyxl.load_workbook(str(output))
    try:
        ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb[wb.sheetnames[0]]
        content = _unwrap_flow_context(context)
        replacements = _context_replacements(content)
        changed = _replace_excel_placeholders(ws, replacements)

        if handling == "preserve":
            if changed == 0:
                raise ValueError("Preserve unchanged requires at least one {{placeholder}} in the Excel template.")
        else:
            header_row = _template_header_row(ws)
            _, logical_width = _excel_meaningful_bounds(ws)
            template_headers = [ws.cell(header_row, col).value for col in range(1, logical_width + 1)]
            rows = parse_excel_rows(content, template_headers=template_headers)
            if start_cell:
                changed += _write_excel_rows_at(
                    ws,
                    rows,
                    start_cell,
                    replace_existing=handling == "replace" and changed == 0,
                )
            elif handling == "replace" and changed == 0:
                changed += _replace_excel_example_rows(ws, rows)
            elif handling == "append" or render_mode == "adaptive":
                changed += _append_excel_rows(ws, rows)
        wb.save(str(output))
        return changed
    finally:
        wb.close()


def _copy_template(source: Path, output: Path) -> None:
    source = source.resolve()
    output = output.resolve()
    if source == output:
        raise ValueError("Template output must be a new file; it cannot overwrite the source template.")
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, output)


def _extract_word_contract(path: Path) -> tuple[str, str]:
    from docx import Document

    doc = Document(str(path))
    headings: list[str] = []
    placeholders: set[str] = set()
    example_parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        example_parts.append(text)
        placeholders.update(_PLACEHOLDER_RE.findall(text))
        if str(paragraph.style.name).lower().startswith("heading"):
            headings.append(text)

    table_shapes: list[str] = []
    for index, table in enumerate(doc.tables, 1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        width = max((len(row) for row in rows), default=0)
        headers = rows[0] if rows else []
        table_shapes.append(f"Table {index}: {len(rows)} row(s) x {width} column(s); header: " + " | ".join(headers))
        for row in rows[:12]:
            example_parts.append("| " + " | ".join(row) + " |")
            for cell in row:
                placeholders.update(_PLACEHOLDER_RE.findall(cell))

    contract_lines = ["Word document structure:"]
    contract_lines.append("- Headings: " + (" > ".join(headings) if headings else "none detected"))
    contract_lines.append("- Placeholders: " + (", ".join(sorted(placeholders)) if placeholders else "none detected"))
    contract_lines.extend(f"- {shape}" for shape in table_shapes)
    contract_lines.append("- Preserve the document's page setup, styles, headers, footers, tables, and branding.")
    return "\n".join(contract_lines), "\n".join(example_parts)


def _word_template_manifest(path: Path) -> dict[str, Any]:
    from docx import Document

    doc = Document(str(path))
    placeholders: set[str] = set()
    for paragraph in doc.paragraphs:
        placeholders.update(_PLACEHOLDER_RE.findall(paragraph.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                placeholders.update(_PLACEHOLDER_RE.findall(cell.text))
    return {
        "kind": "word",
        "placeholders": sorted(placeholders),
        "tables": [
            {
                "index": index,
                "rows": len(table.rows),
                "columns": max((len(row.cells) for row in table.rows), default=0),
                "headers": [cell.text.strip() for cell in table.rows[0].cells] if table.rows else [],
            }
            for index, table in enumerate(doc.tables, 1)
        ],
    }


def _excel_meaningful_bounds(ws: Any) -> tuple[int, int]:
    """Return bounds containing values, excluding formatting-only cells."""
    last_row = 0
    last_column = 0
    for row in ws.iter_rows():
        for cell in row:
            if cell.value not in (None, ""):
                last_row = max(last_row, cell.row)
                last_column = max(last_column, cell.column)
    return max(1, last_row), max(1, last_column)


def _excel_template_manifest(path: Path) -> dict[str, Any]:
    import openpyxl
    from openpyxl.utils import get_column_letter

    wb = openpyxl.load_workbook(str(path), data_only=False, read_only=False)
    try:
        sheets = []
        for ws in wb.worksheets:
            logical_rows, logical_columns = _excel_meaningful_bounds(ws)
            sheets.append(
                {
                    "name": ws.title,
                    "logical_range": f"A1:{get_column_letter(logical_columns)}{logical_rows}",
                    "logical_rows": logical_rows,
                    "logical_columns": logical_columns,
                    "suggested_append_cell": f"A{logical_rows + 1}",
                    "preview_rows": [
                        [_manifest_scalar(ws.cell(row, column).value) for column in range(1, logical_columns + 1)]
                        for row in range(1, min(logical_rows, 3) + 1)
                    ],
                    "tables": [
                        {"name": table.name, "range": table.ref}
                        for table in ws.tables.values()
                    ],
                }
            )
        return {"kind": "excel", "sheets": sheets}
    finally:
        wb.close()


def _manifest_scalar(value: Any) -> Any:
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return str(value)


def _extract_excel_contract(path: Path) -> tuple[str, str]:
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=False, read_only=True)
    try:
        contract_lines = ["Excel workbook structure:"]
        example_parts: list[str] = []
        placeholders: set[str] = set()
        for ws in wb.worksheets:
            logical_rows, logical_columns = _excel_meaningful_bounds(ws)
            rows = list(
                ws.iter_rows(
                    min_row=1,
                    max_row=min(logical_rows, 12),
                    max_col=logical_columns,
                    values_only=True,
                )
            )
            first_nonempty = next((row for row in rows if any(value not in (None, "") for value in row)), ())
            headers = [str(value) if value is not None else "" for value in first_nonempty]
            formula_count = sum(1 for row in rows for value in row if isinstance(value, str) and value.startswith("="))
            contract_lines.append(
                f"- Sheet '{ws.title}': logical range {logical_rows} row(s) x {logical_columns} column(s); "
                f"header: {' | '.join(headers) or 'none detected'}; formulas in preview: {formula_count}"
            )
            example_parts.append(f"[Sheet: {ws.title}]")
            for row in rows:
                values = ["" if value is None else str(value) for value in row]
                example_parts.append("\t".join(values))
                for value in values:
                    placeholders.update(_PLACEHOLDER_RE.findall(value))
        contract_lines.append(
            "- Placeholders: " + (", ".join(sorted(placeholders)) if placeholders else "none detected")
        )
        contract_lines.append(
            "- Preserve sheet names, formulas, formatting, widths, frozen panes, and workbook structure."
        )
        return "\n".join(contract_lines), "\n".join(example_parts)
    finally:
        wb.close()


def _unwrap_flow_context(context: str) -> str:
    text = str(context or "").strip()
    match = re.fullmatch(r"\[Output from '[^']+'\]:\n([\s\S]*)", text)
    return match.group(1).strip() if match else text


def _context_replacements(context: str) -> dict[str, str]:
    replacements = {"context": context}
    candidate = context.strip()
    fenced = re.fullmatch(r"```(?:json)?\s*([\s\S]*?)\s*```", candidate, re.IGNORECASE)
    if fenced:
        candidate = fenced.group(1)
    try:
        parsed = json.loads(candidate)
    except (json.JSONDecodeError, TypeError):
        parsed = None
    if isinstance(parsed, dict):
        source = parsed.get("placeholders") if isinstance(parsed.get("placeholders"), dict) else parsed
        for key, value in source.items():
            if isinstance(value, (str, int, float, bool)) or value is None:
                replacements[str(key)] = "" if value is None else str(value)
    return replacements


def _populate_first_word_table(doc: Any, content: str, append: bool) -> int:
    """Populate a Word table from a structured matrix in adaptive mode."""
    documents = _json_documents(content)
    if documents is None:
        return 0
    if len(documents) == 1 and isinstance(documents[0], dict):
        table_payload = documents[0].get("tables")
        if isinstance(table_payload, dict) and table_payload:
            documents = [next(iter(table_payload.values()))]
        elif isinstance(table_payload, list):
            documents = [table_payload]
    rows = _structured_rows(documents, None)
    if not rows:
        return 0

    if doc.tables:
        table = doc.tables[0]
        width = max((len(row.cells) for row in table.rows), default=max(len(row) for row in rows))
        if not append:
            while len(table.rows) > 1:
                table._tbl.remove(table.rows[-1]._tr)
    else:
        width = max(len(row) for row in rows)
        table = doc.add_table(rows=0, cols=width)

    changed = 0
    for values in rows:
        target = table.add_row()
        for index in range(width):
            target.cells[index].text = str(values[index]) if index < len(values) and values[index] is not None else ""
            changed += 1
    return changed


def _replace_text(text: str, replacements: dict[str, str]) -> tuple[str, int]:
    count = 0

    def replace(match: re.Match[str]) -> str:
        nonlocal count
        key = match.group(1)
        if key not in replacements:
            return match.group(0)
        count += 1
        return replacements[key]

    return _PLACEHOLDER_RE.sub(replace, text), count


def _replace_paragraph(paragraph: Any, replacements: dict[str, str]) -> int:
    updated, count = _replace_text(paragraph.text, replacements)
    if count:
        if paragraph.runs:
            paragraph.runs[0].text = updated
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = updated
    return count


def _replace_docx_placeholders(doc: Any, replacements: dict[str, str]) -> int:
    count = sum(_replace_paragraph(paragraph, replacements) for paragraph in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                count += sum(_replace_paragraph(paragraph, replacements) for paragraph in cell.paragraphs)
    for section in doc.sections:
        for container in (section.header, section.footer):
            count += sum(_replace_paragraph(paragraph, replacements) for paragraph in container.paragraphs)
    return count


def _replace_word_example_content(doc: Any, content: str) -> int:
    lines = content.splitlines() or [content]
    editable = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.strip() and not str(paragraph.style.name).lower().startswith("heading")
    ]
    changed = 0
    for index, paragraph in enumerate(editable):
        paragraph.text = lines[index] if index < len(lines) else ""
        changed += 1
    for line in lines[len(editable) :]:
        doc.add_paragraph(line)
        changed += 1
    return changed


def _replace_excel_placeholders(ws: Any, replacements: dict[str, str]) -> int:
    changed = 0
    for row in ws.iter_rows():
        for cell in row:
            if not isinstance(cell.value, str):
                continue
            updated, count = _replace_text(cell.value, replacements)
            if count:
                cell.value = updated
                changed += count
    return changed


def _canonical_excel_key(value: Any) -> str:
    if value is None:
        return ""
    if hasattr(value, "isoformat") and not isinstance(value, str):
        value = value.isoformat()
    text = str(value).strip()
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}[T ]00:00:00", text):
        text = text[:10]
    return re.sub(r"\s+", " ", text).casefold()


def _decode_json_value(value: Any, depth: int = 0) -> Any:
    """Decode rows that a model or workflow serialized more than once."""
    if not isinstance(value, str) or depth >= 3:
        return value
    candidate = value.strip()
    if not candidate or candidate[0] not in '[{"':
        return value
    try:
        return _decode_json_value(json.loads(candidate), depth + 1)
    except (json.JSONDecodeError, TypeError):
        return value


def _json_documents(context: str) -> list[Any] | None:
    text = context.strip()
    fence_markers = re.findall(r"```", text)
    if fence_markers:
        if len(fence_markers) % 2:
            raise ValueError(
                "The structured Excel payload is incomplete or truncated "
                "(an opening JSON code fence has no closing fence)."
            )
        blocks = re.findall(r"```(?:json)?\s*([\s\S]*?)\s*```", text, re.IGNORECASE)
        if len(blocks) * 2 != len(fence_markers):
            raise ValueError(
                "The structured Excel payload is incomplete or truncated "
                "(an opening JSON code fence has no closing fence)."
            )
        documents: list[Any] = []
        for block in blocks:
            try:
                documents.append(json.loads(block))
            except json.JSONDecodeError as exc:
                raise ValueError(
                    f"The structured Excel payload contains invalid or truncated JSON: {exc.msg} at position {exc.pos}."
                ) from exc
        return documents

    if not text or text[0] not in '[{"':
        return None

    decoder = json.JSONDecoder()
    documents = []
    position = 0
    try:
        while position < len(text):
            while position < len(text) and text[position].isspace():
                position += 1
            if position >= len(text):
                break
            value, position = decoder.raw_decode(text, position)
            documents.append(value)
    except json.JSONDecodeError as exc:
        raise ValueError(
            f"The structured Excel payload contains invalid or truncated JSON: {exc.msg} at position {exc.pos}."
        ) from exc
    return documents


def _object_rows_to_matrix(rows: list[dict[str, Any]], template_headers: list[Any] | None) -> list[list[Any]]:
    if template_headers is None:
        keys: list[str] = []
        for row in rows:
            for key in row:
                if key not in keys:
                    keys.append(key)
        return [keys, *[[row.get(key) for key in keys] for row in rows]]

    canonical_headers = [_canonical_excel_key(header) for header in template_headers]
    nonblank_headers = {header for header in canonical_headers if header}
    canonical_rows = [
        {_canonical_excel_key(key): value for key, value in row.items()}
        for row in rows
    ]
    payload_keys = {key for row in canonical_rows for key in row}
    matched = nonblank_headers.intersection(payload_keys)
    if nonblank_headers and not matched:
        available = sorted(payload_keys)[:12]
        raise ValueError(
            "Structured Excel object keys do not match the selected template headers. "
            f"Payload keys include: {', '.join(available) or 'none'}."
        )
    unmatched = sorted(key for key in payload_keys if key and key not in nonblank_headers)
    if unmatched:
        raise ValueError(
            "Structured Excel object keys are missing from the selected template headers: "
            + ", ".join(unmatched[:12])
            + ("..." if len(unmatched) > 12 else "")
            + "."
        )
    return [[row.get(header) if header else None for header in canonical_headers] for row in canonical_rows]


def _structured_rows(documents: list[Any], template_headers: list[Any] | None) -> list[list[Any]]:
    values: list[Any] = []
    for document in documents:
        document = _decode_json_value(document)
        if isinstance(document, dict):
            for wrapper in ("rows", "data", "items", "records", "structured_output"):
                if wrapper in document and isinstance(_decode_json_value(document[wrapper]), list):
                    document = _decode_json_value(document[wrapper])
                    break
        values.extend(document if isinstance(document, list) else [document])

    values = [_decode_json_value(value) for value in values]
    if not values:
        return []
    if all(isinstance(value, dict) for value in values):
        return _object_rows_to_matrix(values, template_headers)
    if all(isinstance(value, (list, tuple)) for value in values):
        rows = [list(value) for value in values]
        if template_headers is not None:
            expected_width = len(template_headers)
            invalid = [(index + 1, len(row)) for index, row in enumerate(rows) if len(row) != expected_width]
            if invalid:
                sample = ", ".join(f"row {row}: {width}" for row, width in invalid[:8])
                raise ValueError(
                    f"Structured Excel rows must contain exactly {expected_width} values; {sample}."
                )
        return rows
    if all(not isinstance(value, (dict, list, tuple)) for value in values):
        return [[value] for value in values]
    raise ValueError("Structured Excel rows must consistently be arrays or objects; mixed row types were received.")


def parse_excel_rows(context: str, template_headers: list[Any] | None = None) -> list[list[Any]]:
    """Normalize model output into typed worksheet rows.

    Supports JSON arrays/objects (including fenced, chunked, and double-serialized
    output), Markdown tables, and TSV. Invalid structured JSON fails closed so its
    textual representation is never written into worksheet cells.
    """
    content = _unwrap_flow_context(context)
    documents = _json_documents(content)
    if documents is not None:
        return _structured_rows(documents, template_headers)

    lines = content.strip().splitlines() if content.strip() else []
    markdown_rows: list[list[str]] = []
    saw_separator = False
    for line in lines:
        stripped = line.strip()
        if not (stripped.startswith("|") and stripped.endswith("|")):
            continue
        cells = [cell.strip() for cell in stripped[1:-1].split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            saw_separator = True
            continue
        if cells:
            if markdown_rows and cells == markdown_rows[0]:
                continue
            markdown_rows.append(cells)
    if saw_separator and markdown_rows:
        return markdown_rows
    return [line.split("\t") for line in lines]


def _rows_from_context(context: str) -> list[list[Any]]:
    return parse_excel_rows(context)


def _template_header_row(ws: Any) -> int:
    for row_index in range(1, ws.max_row + 1):
        if any(ws.cell(row_index, column).value not in (None, "") for column in range(1, ws.max_column + 1)):
            return row_index
    return 1


def _copy_row_style(ws: Any, source_row: int, target_row: int, width: int) -> None:
    if source_row > ws.max_row:
        return
    for column in range(1, width + 1):
        source = ws.cell(source_row, column)
        target = ws.cell(target_row, column)
        if source.has_style:
            target._style = copy(source._style)
        if source.number_format:
            target.number_format = source.number_format


def _clean_excel_value(value: Any) -> Any:
    return value.strip() if isinstance(value, str) else value


def _replace_excel_example_rows(ws: Any, rows: list[list[Any]]) -> int:
    if not rows:
        return 0
    header_row = _template_header_row(ws)
    existing_end, logical_width = _excel_meaningful_bounds(ws)
    template_headers = [str(ws.cell(header_row, col).value or "").strip() for col in range(1, logical_width + 1)]
    if rows and [_clean_excel_value(cell) for cell in rows[0]][: len(template_headers)] == template_headers:
        rows = rows[1:]
    start_row = header_row + 1
    for row_index in range(start_row, existing_end + 1):
        for column in range(1, logical_width + 1):
            cell = ws.cell(row_index, column)
            if not (isinstance(cell.value, str) and cell.value.startswith("=")):
                cell.value = None
    width = max([logical_width, *(len(row) for row in rows)])
    style_row = start_row if start_row <= existing_end else header_row
    changed = 0
    if rows:
        from openpyxl.utils import get_column_letter

        from shogun.office.adapters.excel_adapter import log_excel_payload_shape

        destination = f"A{start_row}:{get_column_letter(width)}{start_row + len(rows) - 1}"
        log_excel_payload_shape("excel_template_replace", ws.title, destination, rows)
    for offset, values in enumerate(rows):
        target_row = start_row + offset
        _copy_row_style(ws, style_row, target_row, width)
        for column, value in enumerate(values, 1):
            ws.cell(target_row, column, value=_clean_excel_value(value))
            changed += 1
    return changed


def _append_excel_rows(ws: Any, rows: list[list[Any]]) -> int:
    if not rows:
        return 0
    header_row = _template_header_row(ws)
    meaningful_end, logical_width = _excel_meaningful_bounds(ws)
    template_headers = [str(ws.cell(header_row, col).value or "").strip() for col in range(1, logical_width + 1)]
    if [_clean_excel_value(cell) for cell in rows[0]][: len(template_headers)] == template_headers:
        rows = rows[1:]
    start_row = meaningful_end + 1
    width = max([logical_width, *(len(row) for row in rows)], default=logical_width)
    style_row = start_row if start_row <= ws.max_row else meaningful_end
    changed = 0
    if rows:
        from openpyxl.utils import get_column_letter

        from shogun.office.adapters.excel_adapter import log_excel_payload_shape

        destination = f"A{start_row}:{get_column_letter(width)}{start_row + len(rows) - 1}"
        log_excel_payload_shape("excel_template_append", ws.title, destination, rows)
    for offset, values in enumerate(rows):
        target_row = start_row + offset
        _copy_row_style(ws, style_row, target_row, width)
        for column, value in enumerate(values, 1):
            ws.cell(target_row, column, value=_clean_excel_value(value))
            changed += 1
    _expand_excel_tables(ws, start_row, len(rows), width)
    return changed


def _write_excel_rows_at(
    ws: Any,
    rows: list[list[Any]],
    start_cell: str,
    replace_existing: bool = False,
) -> int:
    if not rows:
        return 0
    from openpyxl.utils import coordinate_to_tuple, get_column_letter

    from shogun.office.adapters.excel_adapter import log_excel_payload_shape

    try:
        start_row, start_column = coordinate_to_tuple(start_cell.upper())
    except Exception as exc:
        raise ValueError(f"Invalid template data start cell '{start_cell}'. Use a cell such as A4.") from exc
    width = max(len(row) for row in rows)
    end_column = start_column + width - 1
    if replace_existing:
        existing_end, logical_width = _excel_meaningful_bounds(ws)
        clear_end_column = max(logical_width, end_column)
        for row_index in range(start_row, existing_end + 1):
            for column in range(start_column, clear_end_column + 1):
                cell = ws.cell(row_index, column)
                if not (isinstance(cell.value, str) and cell.value.startswith("=")):
                    cell.value = None
    destination = (
        f"{get_column_letter(start_column)}{start_row}:"
        f"{get_column_letter(end_column)}{start_row + len(rows) - 1}"
    )
    log_excel_payload_shape("excel_template_anchored", ws.title, destination, rows)
    style_row = start_row if start_row <= ws.max_row else max(1, start_row - 1)
    changed = 0
    for row_offset, values in enumerate(rows):
        target_row = start_row + row_offset
        _copy_row_style(ws, style_row, target_row, end_column)
        for column_offset, value in enumerate(values):
            ws.cell(target_row, start_column + column_offset, value=_clean_excel_value(value))
            changed += 1
    _expand_excel_tables(ws, start_row, len(rows), end_column)
    return changed


def _expand_excel_tables(ws: Any, start_row: int, row_count: int, end_column: int) -> None:
    if not row_count:
        return
    from openpyxl.utils import get_column_letter, range_boundaries

    written_end = start_row + row_count - 1
    for table in ws.tables.values():
        min_col, min_row, max_col, max_row = range_boundaries(table.ref)
        if min_col <= 1 <= max_col and start_row <= max_row + 1:
            table.ref = (
                f"{get_column_letter(min_col)}{min_row}:"
                f"{get_column_letter(max(max_col, end_column))}{max(max_row, written_end)}"
            )
